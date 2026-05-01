"""Text extraction from CSV, Excel, PDF, Word, and image files for RAG indexing."""

import base64
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Supported extensions → MIME types accepted at upload
SUPPORTED_EXTENSIONS = {
    ".csv", ".xlsx", ".xls",
    ".pdf",
    ".docx",
    ".jpg", ".jpeg", ".png", ".webp",
    ".txt",
}

MAX_IMAGE_BYTES = 4 * 1024 * 1024   # Groq vision limit: resize if larger
VISION_MODEL = "llama-3.2-11b-vision-preview"


def get_doc_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in (".csv", ".xlsx", ".xls"):
        return "financial"
    if ext in (".jpg", ".jpeg", ".png", ".webp"):
        return "image"
    return "document"


def extract_text(content: bytes, filename: str, groq_client=None) -> str:
    """Dispatch to the correct extractor based on file extension.

    Args:
        content: Raw file bytes.
        filename: Original filename (determines format).
        groq_client: Groq client for image vision extraction. None = skip.

    Returns:
        Extracted plain text, or empty string on failure.
    """
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".csv":
            return _extract_csv(content)
        if ext in (".xlsx", ".xls"):
            return _extract_excel(content)
        if ext == ".pdf":
            return _extract_pdf(content)
        if ext == ".docx":
            return _extract_docx(content)
        if ext in (".jpg", ".jpeg", ".png", ".webp"):
            return _extract_image(content, filename, groq_client)
        if ext == ".txt":
            return content.decode("utf-8", errors="replace")
    except Exception as exc:
        logger.error("Extraction failed for %s: %s", filename, exc)
    return ""


# ── Format extractors ─────────────────────────────────────────────────────────

def _extract_csv(content: bytes) -> str:
    import pandas as pd
    df = pd.read_csv(io.BytesIO(content))
    df.columns = [c.strip().lower() for c in df.columns]
    lines = [f"Columns: {', '.join(df.columns)}"]
    for _, row in df.iterrows():
        lines.append(" | ".join(f"{k}: {v}" for k, v in row.items()))
    return "\n".join(lines)


def _extract_excel(content: bytes) -> str:
    import pandas as pd
    xl = pd.ExcelFile(io.BytesIO(content), engine="openpyxl")
    parts = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        df.columns = [str(c).strip() for c in df.columns]
        parts.append(f"=== Sheet: {sheet} ===")
        parts.append(f"Columns: {', '.join(df.columns)}")
        for _, row in df.iterrows():
            parts.append(" | ".join(f"{k}: {v}" for k, v in row.items()))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(paragraphs)


def _extract_image(content: bytes, filename: str, groq_client=None) -> str:
    """Use Groq vision (LLaMA 3.2 Vision) to describe image content.

    Falls back to a metadata-only string if groq_client is None.
    """
    if groq_client is None:
        logger.warning("No Groq client — image %s will not be analysed", filename)
        return f"Image file: {filename} (vision analysis unavailable)"

    # Resize if above Groq's limit to avoid payload rejection
    image_bytes = _maybe_resize(content)
    b64 = base64.b64encode(image_bytes).decode()
    ext = Path(filename).suffix.lstrip(".").lower()
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"

    try:
        response = groq_client.chat.completions.create(
            model=VISION_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime};base64,{b64}"},
                        },
                        {
                            "type": "text",
                            "text": (
                                "You are a business analyst. Describe everything visible in this image "
                                "with a focus on any numbers, charts, metrics, tables, financial data, "
                                "or business information. Be thorough and specific."
                            ),
                        },
                    ],
                }
            ],
            max_tokens=1024,
        )
        description = response.choices[0].message.content or ""
        logger.info("Vision extraction ok for %s (%d chars)", filename, len(description))
        return f"Image: {filename}\n\n{description}"
    except Exception as exc:
        logger.error("Groq vision failed for %s: %s", filename, exc)
        return f"Image file: {filename} (vision extraction failed)"


def _maybe_resize(content: bytes) -> bytes:
    """Resize image to stay under MAX_IMAGE_BYTES using Pillow."""
    if len(content) <= MAX_IMAGE_BYTES:
        return content
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        # Scale down proportionally until under limit
        quality = 85
        while quality >= 40:
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=quality)
            if buf.tell() <= MAX_IMAGE_BYTES:
                return buf.getvalue()
            quality -= 15
        # Last resort: halve dimensions
        img = img.resize((img.width // 2, img.height // 2))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=75)
        return buf.getvalue()
    except Exception:
        return content[:MAX_IMAGE_BYTES]
